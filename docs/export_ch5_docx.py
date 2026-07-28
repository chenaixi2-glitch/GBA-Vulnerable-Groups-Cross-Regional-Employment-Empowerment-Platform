#!/usr/bin/env python3
"""将终期报告第5章 Markdown 导出为 Word (.docx)。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "终期报告_第5章_系统设计与技术实现.md"
OUT_PATH = ROOT / "终期报告_第5章_系统设计与技术实现_洪扬萱.docx"


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = row[j] if j < len(row) else ""


def convert(md_text: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] | None = None
    skip_until = None

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if skip_until:
            if line.strip() == skip_until:
                skip_until = None
            continue

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Intense Quote"
                code_lines = []
                in_code = False
            else:
                if line.strip().startswith("```mermaid"):
                    skip_until = "```"
                    continue
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            if not table_rows:
                table_rows = []
            if re.match(r"^\|\s*:?-+", line):
                continue
            cells = [strip_md_inline(c) for c in line.strip("|").split("|")]
            table_rows.append(cells)
            continue
        elif table_rows:
            add_table(doc, table_rows)
            table_rows = None

        if not line.strip():
            continue
        if line.strip() == "---":
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(strip_md_inline(line[2:]), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md_inline(line[3:]), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md_inline(line[4:]), level=2)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(strip_md_inline(line[2:]))
            p.paragraph_format.left_indent = Pt(18)
            continue
        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(strip_md_inline(line[2:]), style="List Bullet")
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            doc.add_paragraph(strip_md_inline(m.group(2)), style="List Number")
            continue

        p = doc.add_paragraph(strip_md_inline(line))

    if table_rows:
        add_table(doc, table_rows)

    return doc


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    cn = len(re.findall(r"[\u4e00-\u9fff]", text))
    print(f"源文件汉字约: {cn}")

    doc = convert(text)
    doc.save(OUT_PATH)
    print(f"已导出: {OUT_PATH}")


if __name__ == "__main__":
    main()
