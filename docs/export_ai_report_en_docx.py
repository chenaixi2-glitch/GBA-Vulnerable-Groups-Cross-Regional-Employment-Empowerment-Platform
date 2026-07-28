#!/usr/bin/env python3
"""Export the English GBA-VEEP report (Ch.5, 11, 12) to a formatted Word document."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Please install: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "AI_Report_EN.md"
OUT_PATH = ROOT / "AI_Report_EN.docx"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
CODE_BG = "F2F4F7"
CODE_COLOR = RGBColor(0x1E, 0x29, 0x3B)

# Inline tokens that should render as code even without backticks in source MD
CODE_TOKEN_RE = re.compile(
    r"(`[^`]+`"
    r"|[A-Za-z_][\w./-]*\.(?:js|py|html|sql|yml|yaml|conf|md)"
    r"|/(?:api|mcp|health|individual|corporate)(?:/[^\s,;.)]+)?"
    r"|\b(?:GET|POST|PUT|PATCH|DELETE)\s+/[^\s,;]+"
    r"|\b[A-Za-z_][\w]*\(\)"
    r"|\b(?:gba_website|ai_career(?:_copilot)?|group_types|target_criteria|"
    r"vulnerable_group_friendly|interview_format|invite_token|match_score|"
    r"match_reasons|resume_snapshot|content_json|skills_text|source_url|"
    r"external_id|is_active_on_source|has_access|has_premium_access|"
    r"disability_type|career_gap_years|current_income|org_invite_code|"
    r"invited_by_user_id|overall_score|category_scores|debrief_summary|"
    r"job_postings|user_resumes|job_applications|job_match_impressions|"
    r"job_external_interests|legal_aid_requests|legal_aid_responses|"
    r"interview_invites|company_orgs|company_org_members|company_profiles|"
    r"JWT_SECRET|REDIS_HOST|MYSQL_\*|SILICONFLOW_API_KEY|DASHSCOPE_API_KEY"
    r"|inferGroupTypes|getPlatformAccess|userMatchesJobCriteria|scoreJobResume|"
    r"sortApplicantsForCorporate|listMatchedForUser|assertInternalJobOwner|"
    r"upsert_external_job|authenticate|requireRole|localStorage|host\.docker\.internal"
    r")\b"
    r"|\*\*[^*]+\*\*)"
)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _shade_run(run, fill: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    r_pr.append(shd)


def _apply_body_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for level, size, color in (
        (0, 20, RGBColor(0x0F, 0x17, 0x2A)),
        (1, 18, RGBColor(0x1A, 0x36, 0x5D)),
        (2, 14, RGBColor(0x1E, 0x40, 0x6E)),
        (3, 12, RGBColor(0x25, 0x63, 0xEB)),
    ):
        style_name = "Title" if level == 0 else f"Heading {level}"
        h = doc.styles[style_name]
        h.font.name = BODY_FONT
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18 if level <= 1 else 12)
        h.paragraph_format.space_after = Pt(8 if level <= 1 else 6)


def _add_code_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = CODE_COLOR
    _shade_run(run, CODE_BG)


def _add_text_run(paragraph, text: str, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.bold = bold


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in CODE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            chunk = text[pos : m.start()]
            if chunk:
                _add_text_run(p, chunk)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_text_run(p, token[2:-2], bold=True)
        elif token.startswith("`") and token.endswith("`"):
            _add_code_run(p, token[1:-1])
        else:
            _add_code_run(p, token)
        pos = m.end()
    if pos < len(text):
        _add_text_run(p, text[pos:])


def strip_md_inline(text: str) -> str:
    return text.strip()


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_paragraph_in_place(p, cell_text)
            if header and i == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_shading(cell, "E8EEF7")
    doc.add_paragraph("")


def add_rich_paragraph_in_place(paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("r"):
            element.remove(child)
    pos = 0
    for m in CODE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            chunk = text[pos : m.start()]
            if chunk:
                _add_text_run(paragraph, chunk)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_text_run(paragraph, token[2:-2], bold=True)
        elif token.startswith("`") and token.endswith("`"):
            _add_code_run(paragraph, token[1:-1])
        else:
            _add_code_run(paragraph, token)
        pos = m.end()
    if pos < len(text):
        _add_text_run(paragraph, text[pos:])


def heading_level(line: str) -> int | None:
    s = line.strip()
    if s.startswith("# "):
        return 0
    if s.startswith("## "):
        return 1
    if s.startswith("### "):
        return 2
    if s.startswith("#### "):
        return 3
    m = re.match(r"^(\d+(?:\.\d+)*)\s+\S", s)
    if not m:
        return None
    depth = m.group(1).count(".")
    if depth == 0:
        return 0
    if depth == 1:
        return 1
    if depth == 2:
        return 2
    return 3


def convert(md_text: str) -> Document:
    doc = Document()
    _apply_body_style(doc)

    sections = doc.sections[0]
    sections.top_margin = Inches(1)
    sections.bottom_margin = Inches(1)
    sections.left_margin = Inches(1.1)
    sections.right_margin = Inches(1.1)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] | None = None
    skip_until: str | None = None

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if skip_until:
            if line.strip() == skip_until:
                skip_until = None
            continue

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = CODE_FONT
                run.font.size = Pt(9.5)
                run.font.color.rgb = CODE_COLOR
                p.paragraph_format.left_indent = Inches(0.2)
                _shade_run(run, CODE_BG)
                code_lines = []
                in_code = False
            else:
                if line.strip().startswith("```mermaid"):
                    skip_until = "```"
                else:
                    in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            if table_rows is None:
                table_rows = []
            if re.match(r"^\|\s*:?-+", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_rows.append(cells)
            continue
        if table_rows:
            add_table(doc, table_rows)
            table_rows = None

        if not line.strip():
            continue
        if line.strip() == "---":
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(strip_md_inline(line[2:]), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md_inline(line[3:]), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md_inline(line[4:]), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(strip_md_inline(line[5:]), level=3)
            continue

        hl = heading_level(line)
        if hl is not None and re.match(r"^\d+(?:\.\d+)+\s+", line.strip()):
            doc.add_heading(strip_md_inline(line.strip()), level=min(hl + 1, 3))
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_rich_paragraph_in_place(p, line[2:])
            continue

        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_paragraph_in_place(p, line[2:])
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich_paragraph_in_place(p, m.group(2))
            continue

        if re.match(r"^（\d+）", line) or re.match(r"^\(\d+\)", line):
            p = doc.add_paragraph()
            add_rich_paragraph_in_place(p, line)
            p.paragraph_format.space_before = Pt(4)
            continue

        add_rich_paragraph(doc, line)

    if table_rows:
        add_table(doc, table_rows)

    return doc


def main() -> None:
    if not MD_PATH.exists():
        print(f"Missing source: {MD_PATH}")
        sys.exit(1)
    text = MD_PATH.read_text(encoding="utf-8")
    doc = convert(text)
    doc.save(OUT_PATH)
    print(f"Exported: {OUT_PATH}")


if __name__ == "__main__":
    main()
