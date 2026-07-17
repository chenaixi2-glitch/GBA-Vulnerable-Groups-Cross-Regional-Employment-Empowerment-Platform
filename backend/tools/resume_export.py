"""Resume export helpers — HTML → PDF, ResumeContent → DOCX."""

from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import TYPE_CHECKING
from urllib.parse import quote

from tools.resume_layout import SECTION_LABELS, normalize_language
from log import get_logger

logger = get_logger("app")

if TYPE_CHECKING:
    from workflow.state import ResumeContent

_TEMPLATE_DIR = Path(__file__).parent.parent / "templates"
_MSYS2_DLL_DIRS = (
    Path(r"C:\msys64\mingw64\bin"),
    Path(r"C:\msys64\ucrt64\bin"),
)

_PDF_FONT_STACK = (
    "'Noto Sans CJK SC', 'Source Han Sans SC', 'Microsoft YaHei', "
    "'PingFang SC', 'Hiragino Sans GB', 'WenQuanYi Micro Hei', "
    "'Inter', 'Segoe UI', sans-serif"
)

_PDF_INJECTED_CSS = f"""
/* WeasyPrint PDF overrides */
@page {{
    size: A4 portrait;
    margin: 0;
}}
html, body {{
    background: #fff !important;
    color: #111 !important;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
}}
html.dark body,
html.dark .section h2,
html.dark .item-content {{
    background: #fff !important;
    color: #111 !important;
    border-color: #333 !important;
}}
body, .resume-container {{
    font-family: {_PDF_FONT_STACK} !important;
}}
.resume-container {{
    width: 210mm !important;
    min-height: auto !important;
    max-height: none !important;
    overflow: visible !important;
    margin: 0 auto !important;
}}
.section {{
    page-break-inside: avoid;
}}
"""


class WeasyPrintUnavailableError(RuntimeError):
    """Raised when WeasyPrint or its native dependencies are missing."""

    INSTALL_HINT = (
        "PDF 导出依赖 WeasyPrint 未就绪。"
        "Linux/Docker: sudo bash scripts/install-weasyprint-linux.sh；"
        "Windows: .\\scripts\\install-weasyprint-windows.ps1"
    )

    def __init__(self, cause: Exception | None = None) -> None:
        super().__init__(self.INSTALL_HINT)
        self.cause = cause


def _ensure_weasyprint_dll_path() -> None:
    """On Windows, point WeasyPrint to MSYS2 Pango DLLs when env is unset."""
    if os.name != "nt" or os.environ.get("WEASYPRINT_DLL_DIRECTORIES"):
        return
    for dll_dir in _MSYS2_DLL_DIRS:
        if (dll_dir / "libpango-1.0-0.dll").exists():
            os.environ["WEASYPRINT_DLL_DIRECTORIES"] = str(dll_dir)
            return


def weasyprint_available() -> bool:
    """Return True if WeasyPrint and native libraries can be loaded."""
    try:
        _ensure_weasyprint_dll_path()
        from weasyprint import HTML  # noqa: F401
        return True
    except OSError:
        return False


def sanitize_export_filename(name: str | None, ext: str, fallback: str = "resume") -> str:
    """Build a safe attachment filename, preserving CJK when possible."""
    raw = (name or "").strip() or fallback
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", raw).strip().strip(".")
    base = cleaned[:80] if cleaned else fallback
    ext_clean = re.sub(r"[^a-z0-9]", "", ext.lower()) or "pdf"
    return f"{base}.{ext_clean}"


def build_content_disposition(filename: str) -> str:
    """RFC 5987 Content-Disposition with UTF-8 filename* fallback."""
    ascii_name = filename.encode("ascii", "ignore").decode("ascii") or "resume.pdf"
    encoded = quote(filename, safe="")
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{encoded}'


def prepare_html_for_pdf(html: str) -> str:
    """Normalize resume HTML for reliable WeasyPrint rendering."""
    content = (html or "").strip()
    if not content:
        raise ValueError("简历 HTML 为空，无法导出 PDF")

    lower = content.lower()
    if "<html" not in lower:
        content = (
            f'<!DOCTYPE html><html lang="zh"><head><meta charset="UTF-8">'
            f"<title>Resume</title></head><body>{content}</body></html>"
        )
        lower = content.lower()

    if _PDF_INJECTED_CSS.strip() not in content:
        if "</head>" in lower:
            idx = lower.rindex("</head>")
            content = content[:idx] + f"<style>{_PDF_INJECTED_CSS}</style>" + content[idx:]
        else:
            content = f"<style>{_PDF_INJECTED_CSS}</style>{content}"

    # PDF should always use light theme for readability
    content = re.sub(r'\bclass="([^"]*\b)dark(\b[^"]*)"', r'class="\1\2"', content)
    content = content.replace(' class="dark"', "").replace('class="dark"', "")

    return content


def html_to_pdf_bytes(html: str) -> bytes:
    """Convert resume HTML to PDF bytes via WeasyPrint."""
    prepared = prepare_html_for_pdf(html)
    _ensure_weasyprint_dll_path()
    try:
        from weasyprint import HTML
    except OSError as exc:
        raise WeasyPrintUnavailableError(exc) from exc

    try:
        return HTML(string=prepared, base_url=str(_TEMPLATE_DIR)).write_pdf()
    except OSError as exc:
        raise WeasyPrintUnavailableError(exc) from exc


def count_pdf_pages_from_html(html: str) -> int | None:
    """Render HTML with WeasyPrint and return PDF page count, or None if unavailable."""
    if not weasyprint_available():
        return None
    prepared = prepare_html_for_pdf(html)
    _ensure_weasyprint_dll_path()
    try:
        from weasyprint import HTML
    except OSError:
        return None
    try:
        document = HTML(string=prepared, base_url=str(_TEMPLATE_DIR)).render()
        return len(document.pages)
    except OSError as exc:
        logger.warning("PDF page count failed: %s", exc)
        return None


# Match PDF body text (#111); Word Heading styles default to theme blue.
_DOCX_TEXT_COLOR = "111111"
_DOCX_INLINE_TAGS = frozenset({"a", "strong", "b", "em", "i", "u", "span", "br", "code", "font", "small"})
_DOCX_BLOCK_TAGS = frozenset({"div", "section", "header", "footer", "article", "main", "aside"})


def _clear_color_theme_attrs(color_el) -> None:
    from docx.oxml.ns import qn

    color_el.set(qn("w:val"), _DOCX_TEXT_COLOR)
    for attr in ("themeColor", "themeTint", "themeShade"):
        key = qn(f"w:{attr}")
        if key in color_el.attrib:
            del color_el.attrib[key]


def _prepare_html_for_docx(html: str) -> str:
    """Flatten resume HTML so HtmlToDocx does not glue body text into blue Heading runs.

    HtmlToDocx keeps the current paragraph after </h2>, and resume templates put
    education / experience in <div>s — so body text inherits Heading (theme blue).
    Convert headings and leaf blocks into explicit <p> tags first.
    """
    from bs4 import BeautifulSoup, NavigableString

    soup = BeautifulSoup(html or "", "html.parser")
    for tag in soup.find_all(["style", "script", "img"]):
        tag.decompose()

    def _phrasing_only(tag) -> bool:
        for child in tag.children:
            if isinstance(child, NavigableString):
                continue
            name = getattr(child, "name", None)
            if name in _DOCX_INLINE_TAGS and _phrasing_only(child):
                continue
            return False
        return True

    for level in range(1, 7):
        for heading in soup.find_all(f"h{level}"):
            paragraph = soup.new_tag("p")
            strong = soup.new_tag("strong")
            strong.string = heading.get_text()
            paragraph.append(strong)
            heading.replace_with(paragraph)

    for tag in reversed(soup.find_all(list(_DOCX_BLOCK_TAGS | {"span"}))):
        if tag.name is None:
            continue
        if not tag.get_text(strip=True):
            tag.unwrap()
            continue
        if _phrasing_only(tag):
            tag.name = "p"

    body = soup.body or soup
    return str(body)


def _normalize_docx_text_color(doc) -> None:
    """Force resume DOCX text to near-black so it matches PDF (not Word theme blue)."""
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    rgb = RGBColor(0x11, 0x11, 0x11)

    for style in doc.styles:
        try:
            if style.font is not None:
                style.font.color.rgb = rgb
        except (AttributeError, ValueError, KeyError, TypeError):
            pass
        element = getattr(style, "element", None)
        if element is not None:
            for color_el in element.iter(qn("w:color")):
                _clear_color_theme_attrs(color_el)

    for color_el in doc.element.iter(qn("w:color")):
        _clear_color_theme_attrs(color_el)

    def _paint_paragraphs(paragraphs) -> None:
        for para in paragraphs:
            # Never leave body content on Word Heading styles (theme accent = blue).
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                para.style = doc.styles["Normal"]
            for run in para.runs:
                run.font.color.rgb = rgb

    _paint_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _paint_paragraphs(cell.paragraphs)


def _add_docx_section_title(doc, text: str) -> None:
    """Add a section title as bold Normal text (avoid Word Heading theme blue)."""
    from docx.shared import Pt, RGBColor

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)


def html_to_docx_bytes(html: str) -> bytes:
    """Convert rendered resume HTML to a Word document."""
    from docx import Document
    from htmldocx import HtmlToDocx

    if not html or not html.strip():
        raise ValueError("Empty HTML")

    prepared = _prepare_html_for_docx(html)
    doc = Document()
    HtmlToDocx().add_html_to_document(prepared, doc)
    _normalize_docx_text_color(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def resume_content_to_docx_bytes(content: "ResumeContent") -> bytes:
    """Build a Word document from structured resume content."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    lang = normalize_language(content.meta.language)
    labels = SECTION_LABELS.get(lang, SECTION_LABELS["zh"])
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Source Han Sans"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    profile = content.profile
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile.name or "Resume")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    contact_parts = [
        profile.email,
        profile.phone,
        profile.city,
        profile.github,
        profile.linkedin,
        profile.address,
    ]
    contact = " | ".join(part for part in contact_parts if part)
    if contact:
        doc.add_paragraph(contact)

    for edu in profile.education:
        line = " · ".join(
            part for part in [edu.school, edu.major, edu.degree, f"{edu.start_date} - {edu.end_date}"] if part
        )
        if line:
            doc.add_paragraph(line)

    if content.summary:
        _add_docx_section_title(doc, labels["summary"])
        doc.add_paragraph(content.summary)

    sections = [
        ("skills", content.skills),
        ("projects", content.projects),
        ("internships", content.internships),
        ("awards", content.awards),
        ("papers", content.papers),
    ]
    for key, items in sections:
        if not items:
            continue
        _add_docx_section_title(doc, labels.get(key, key))
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run(item.title)
            run.bold = True
            if item.content:
                doc.add_paragraph(item.content)

    _normalize_docx_text_color(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
