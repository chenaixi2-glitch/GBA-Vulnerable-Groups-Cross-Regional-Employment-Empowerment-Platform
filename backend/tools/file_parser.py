"""文件解析工具 — 支持 PDF / DOCX / Markdown / TXT。"""

from __future__ import annotations

import io
from pathlib import Path

from log import get_logger

logger = get_logger("app")

_SUPPORTED_UPLOAD_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# 解析4种文件：pdf、docx、md、txt。对于API上传的内容，支持根据filename后缀解析bytes数据。
def parse_file(file_path: str | Path) -> str:
    """解析文件内容为纯文本/Markdown。"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix in (".md", ".txt"):
        return _parse_text(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def parse_content(content: str, filename: str = "") -> str:
    """解析内容字符串（用于 API 上传的 base64 解码后内容）。"""
    return parse_content_bytes(content.encode("utf-8"), filename=filename)


def parse_content_bytes(data: bytes, filename: str = "") -> str:
    """解析内存中的文件内容。"""
    suffix = Path(filename).suffix.lower() if filename else ""
    if suffix == ".pdf":
        return _parse_pdf_bytes(data)
    if suffix == ".docx":
        return _parse_docx_bytes(data)
    if suffix in (".md", ".txt", ""):
        return data.decode("utf-8")
    raise ValueError(f"Unsupported file type: {suffix}")


def supported_upload_suffixes() -> set[str]:
    """返回允许上传并解析的文件扩展名。"""
    return set(_SUPPORTED_UPLOAD_SUFFIXES)


def _parse_pdf(path: Path) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        result = "\n\n".join(text_parts)
        logger.info("Parsed PDF: %s (%d chars)", path.name, len(result))
        return result
    except ImportError:
        logger.error("pdfplumber not installed")
        raise
    except Exception as e:
        logger.error("Failed to parse PDF %s: %s", path, e)
        raise


def _parse_pdf_bytes(data: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        result = "\n\n".join(text_parts)
        logger.info("Parsed DOCX: %s (%d chars)", path.name, len(result))
        return result
    except ImportError:
        logger.error("python-docx not installed")
        raise
    except Exception as e:
        logger.error("Failed to parse DOCX %s: %s", path, e)
        raise


def _parse_docx_bytes(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(text_parts)


def _parse_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    logger.info("Parsed text: %s (%d chars)", path.name, len(text))
    return text
